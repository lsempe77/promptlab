"""Generate a .docx email introducing the 3ie Prompt Lab tool and its rollout plan."""

from docx import Document
from docx.shared import Pt
from pathlib import Path

doc = Document()
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# Subject
p = doc.add_paragraph()
run = p.add_run("Subject: Agentic 3ie Prompt Lab — a self-improving LLM extraction/screening tool for DEP and beyond")
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph("Hi all,")
doc.add_paragraph(
    "I wanted to share a tool I've been building called the Agentic 3ie Prompt Lab — a platform "
    "that actively tests prompts for screening and extraction itself, judges its own mistakes, "
    "and rewrites its own instructions to get better over time, without anyone having to sit and "
    "babysit it. It's live now as a pilot on a small sample of the DEP (the beauty of the central "
    "limit theorem), and I think it has clear applications across several of our other projects, "
    "so I wanted to explain what it does and where it's headed."
)

# --- What it is ---
doc.add_heading("What makes it \u201cagentic\u201d", level=2)
doc.add_paragraph(
    "Most uses of AI in our work involve a person picking a model, writing a prompt, and "
    "manually checking whether the output looks right. This tool works differently: it acts as "
    "its own agent. On a continuous loop, it (1) runs a task across many different models side "
    "by side, (2) scores every answer against a ground-truth reference set using three "
    "complementary accuracy checks, (3) identifies exactly where and why each model is getting "
    "things wrong, and (4) automatically rewrites its own prompt to fix those specific mistakes "
    "— then re-tests the new version to see if it actually helped. Nobody needs to manually "
    "tweak a prompt and re-test it by hand; the system does that whole cycle by itself, over and "
    "over. Everything — model comparisons, accuracy breakdowns, confusion matrices, and the full "
    "history of every prompt rewrite it made and why — is visible in a live dashboard, so the "
    "process stays fully transparent rather than a black box."
)
doc.add_paragraph(
    "It runs continuously in the cloud (not on my laptop), so it keeps testing, judging, and "
    "improving itself around the clock with no manual intervention. Each project runs as its own "
    "workspace inside the same platform, so results stay directly comparable across them."
)
p = doc.add_paragraph()
p.add_run("Live dashboard: ").bold = True
p.add_run("https://sempe.dev/promptlab")
p = doc.add_paragraph()
p.add_run("GitHub repo: ").bold = True
p.add_run("https://github.com/lsempe77/promptlab")

# --- How it measures quality ---
doc.add_heading("How it decides what's \u201ccorrect\u201d \u2014 and whether a model can be trusted", level=2)
doc.add_paragraph(
    "A single \u201c% correct\u201d number can be misleading, so the tool scores every answer several "
    "ways and looks hard at how a model fails, not just how often:"
)
_quality_bullets = [
    ("Three independent accuracy checks",
     "a fuzzy string match, an exact match, and a separate LLM \u201cjudge\u201d that reads the answer "
     "in context. The judge is deliberately cross-family (e.g. an OpenAI answer is graded by an "
     "Anthropic model, and vice versa) so no model can flatter its own family."),
    ("Honesty, not just accuracy",
     "it separates a model that correctly says \u201cnot stated\u201d from one that confidently makes "
     "something up. A confident wrong answer (a hallucination) is penalised more heavily than an "
     "honest \u201cI don't know\u201d, and each model's hallucination, wrong-answer, and abstention rates "
     "are tracked separately."),
    ("Evidence checking",
     "for each value it extracts, the model has to quote the supporting text, and the tool "
     "verifies that quote actually appears in the source document \u2014 catching invented evidence."),
    ("Confidence calibration",
     "it compares how confident a model claims to be with how often it is actually right (a "
     "reliability diagram and a Brier score), so we know whether a model's confidence can be "
     "taken at face value."),
    ("Consistency",
     "it measures how often different models agree with each other, and how often a single model "
     "gives the same answer when asked the same question repeatedly."),
]
for _name, _desc in _quality_bullets:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(_name).bold = True
    p.add_run(f" \u2014 {_desc}")

# --- When a result is trustworthy ---
doc.add_heading("Knowing when a result is solid enough to rely on", level=2)
doc.add_paragraph(
    "Rather than reading too much into a number from a handful of documents, the tool rolls out "
    "in stages (30 \u2192 60 \u2192 100 reference documents) and reports a 95% confidence interval around "
    "each accuracy that visibly narrows as the sample grows \u2014 so we can tell a real difference "
    "between models from noise. Each model-and-field combination then has to clear a quality gate "
    "(currently \u226580% on the independent judge) before it counts as production-ready, and the "
    "dashboard shows, field by field, how many models currently pass. When no model can clear a "
    "field, that usually points to noise in the ground-truth reference data itself \u2014 a useful "
    "finding in its own right, not just a model failure."
)

# --- Where it's being used ---
doc.add_heading("Where it's being used", level=2)
rows = [
    ("DEP", "Pilot use case: structured-field extraction (sector, authors, affiliations, "
            "countries, sub-sector) from previously extracted .MDs (via Apache Tika, by the DIG "
            "team — thanks!). Screening (title/abstract relevance) is the planned next step once "
            "extraction is running smoothly."),
    ("HSF", "Planned: screening first, with extraction to follow once the screening workflow "
            "is validated."),
    ("Girl Effect", "Planned: screening first, with extraction to follow later."),
    ("StrongMinds", "Planned: screening first, with extraction to follow later."),
]
for name, desc in rows:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(name)
    run.bold = True
    p.add_run(f" — {desc}")
doc.add_paragraph(
    "The idea is that DEP acts as the proving ground for the agent — whatever prompts, models, "
    "and scoring approach it converges on there carry over directly to the other projects, so "
    "each new project starts from an already self-tuned baseline instead of from scratch."
)

# --- Cost/quality frontier ---
doc.add_heading("Finding the cost/quality frontier", level=2)
doc.add_paragraph(
    "I'm currently letting the agent benchmark 15 models across free, cheap, mid-tier, and "
    "premium options (OpenAI, Anthropic, Qwen, Z-ai, Mistral, Gemini, Deepseek, Llama, Gemma, "
    "Poolside). That breadth is deliberate at this stage — it lets us see empirically which "
    "models are actually worth paying for on this kind of task, rather than assuming the most "
    "expensive model is always best. It will cost a bit of money (~US$1,000, which I plan to "
    "split across the four projects). Once we have enough data, we'll drop the models that are "
    "consistently weak or clearly dominated (lower accuracy at a higher price) and settle on a "
    "shortlist that sits on the cost/quality frontier. That should meaningfully reduce the cost "
    "of running this at scale across projects without sacrificing accuracy."
)

# --- Possible paper ---
doc.add_heading("A possible write-up", level=2)
doc.add_paragraph(
    "Once we have solid results — model comparisons, how much the self-rewriting prompts "
    "actually improved over the starting point, and the cost/quality trade-offs — I think it "
    "could be worth writing up as a short methods paper or note. A transparent, agentic approach "
    "to LLM extraction/screening — one that measures honesty and calibration, not just raw "
    "accuracy — doesn't seem to be well documented elsewhere for evaluation research, and it "
    "could be useful for other organisations facing the same problem. Happy to sketch an outline "
    "if there's interest."
)

doc.add_paragraph("")
doc.add_paragraph(
    "Let me know if you'd like a walkthrough of the dashboard or would like to discuss how this "
    "could fit into your project's workflow. In the meantime you can watch the results build up "
    "live — I'm deliberately running just one cloud machine to keep costs down, so it progresses "
    "a bit slowly."
)
doc.add_paragraph("")
doc.add_paragraph("Best,")
doc.add_paragraph("Lucas")

out = Path(
    r"C:\Users\LucasSempe\OneDrive - International Initiative for Impact Evaluation\Desktop\DEP\promptlab_intro_email.docx"
)
doc.save(out)
print(f"Saved: {out}")
